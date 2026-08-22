"""
Genera el reporte de riesgo como documento Word (.docx) con formato
profesional, usando la misma paleta e identidad visual de la app
("expediente de seguridad"): Cambria para títulos, Calibri para texto,
Consolas para cifras/datos, acento petróleo y colores semánticos de
riesgo (verde/ámbar/rojo).

No depende de Flask ni de nada web: recibe los mismos dicts que ya arma
server.py (evaluacion, prediccion, plan) y devuelve un BytesIO listo para
enviar como descarga.
"""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

FUENTE_TITULO = "Cambria"
FUENTE_TEXTO = "Calibri"
FUENTE_DATOS = "Consolas"

INK = RGBColor(0x10, 0x1B, 0x1A)
INK_SOFT = RGBColor(0x52, 0x62, 0x5F)
INK_FAINT = RGBColor(0x8A, 0x97, 0x93)
ACCENT = RGBColor(0x0E, 0x6B, 0x66)
GOOD = RGBColor(0x1F, 0x8A, 0x56)
WARN = RGBColor(0xC2, 0x76, 0x0F)
CRITICAL = RGBColor(0xC2, 0x3B, 0x2E)
LINE_HEX = "D6DEDA"
PAPER_SUNKEN_HEX = "E4E9E6"
ACCENT_HEX = "0E6B66"

COLOR_RIESGO = {"Bajo": GOOD, "Medio": WARN, "Alto": CRITICAL}

NOMBRES_CATEGORIA = {
    "contraseñas": "Contraseñas", "actualizaciones": "Actualizaciones", "redes": "Redes",
    "respaldo": "Respaldo", "dispositivo": "Dispositivo",
}
ORDEN_CATEGORIAS = ["contraseñas", "actualizaciones", "dispositivo", "redes", "respaldo"]
ETIQUETA_SEVERIDAD = {"alto": "Alto", "medio": "Medio", "bajo": "Bajo"}
ETIQUETA_ESFUERZO = {"bajo": "esfuerzo bajo", "medio": "esfuerzo medio", "alto": "esfuerzo alto"}


def _etiquetar_riesgo(puntaje: int) -> str:
    if puntaje < 35:
        return "Bajo"
    if puntaje < 65:
        return "Medio"
    return "Alto"


def _set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=120, right=120) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for nombre, valor in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{nombre}")
        el.set(qn("w:w"), str(valor))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _add_bottom_border(paragraph, hex_color: str, size: int = 16) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _heading(doc: Document, texto: str, color: RGBColor = ACCENT, size: int = 15) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(texto.upper())
    run.font.name = FUENTE_TITULO
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    _add_bottom_border(p, LINE_HEX, size=10)


def _kv_table(doc: Document, filas: list[tuple[str, str, RGBColor | None]]) -> None:
    tabla = doc.add_table(rows=0, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT
    tabla.autofit = False
    for etiqueta, valor, color_valor in filas:
        row = tabla.add_row()
        c1, c2 = row.cells
        c1.width = Inches(2.3)
        c2.width = Inches(3.7)
        _set_cell_shading(c1, PAPER_SUNKEN_HEX)
        _set_cell_margins(c1)
        _set_cell_margins(c2)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(etiqueta)
        r1.font.name = FUENTE_TEXTO
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = INK_SOFT
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(valor)
        r2.font.name = FUENTE_DATOS
        r2.font.size = Pt(11)
        r2.font.bold = True
        r2.font.color.rgb = color_valor or INK
    return tabla


def generar_reporte_docx(evaluacion: dict, prediccion: dict, plan: list[dict]) -> io.BytesIO:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = FUENTE_TEXTO
    estilo_normal.font.size = Pt(10.5)
    estilo_normal.font.color.rgb = INK

    # --- Portada / encabezado -------------------------------------------------
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_after = Pt(2)
    r = p_titulo.add_run("Mapa de Riesgos Digitales")
    r.font.name = FUENTE_TITULO
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = INK

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(2)
    r = p_sub.add_run("EXPEDIENTE DE SEGURIDAD PERSONAL")
    r.font.name = FUENTE_TEXTO
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    p_fecha = doc.add_paragraph()
    p_fecha.paragraph_format.space_after = Pt(0)
    r = p_fecha.add_run("Generado el " + datetime.now().strftime("%d/%m/%Y a las %H:%M"))
    r.font.name = FUENTE_TEXTO
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = INK_FAINT
    _add_bottom_border(p_fecha, ACCENT_HEX, size=20)

    # --- Resumen ----------------------------------------------------------
    _heading(doc, "Resumen")

    categoria_ml = prediccion["categoria_predicha"]
    confianza = round(prediccion["probabilidades"].get(categoria_ml, 0) * 100)
    puntaje = evaluacion["puntaje_global_reglas"]
    categoria_reglas = _etiquetar_riesgo(puntaje)

    _kv_table(doc, [
        ("Puntaje de riesgo (reglas)", f"{puntaje} / 100", COLOR_RIESGO[categoria_reglas]),
        ("Riesgo global (Modelo ML)", f"{categoria_ml}  ·  {confianza}% confianza", COLOR_RIESGO.get(categoria_ml, INK)),
    ])

    # --- Puntaje por categoría ---------------------------------------------
    _heading(doc, "Puntaje por categoría")
    filas_cat = []
    for cat in ORDEN_CATEGORIAS:
        valor = evaluacion["puntajes_por_categoria"].get(cat, 0)
        filas_cat.append((NOMBRES_CATEGORIA[cat], f"{valor} / 100", COLOR_RIESGO[_etiquetar_riesgo(valor)]))
    _kv_table(doc, filas_cat)

    # --- Plan de acción -----------------------------------------------------
    _heading(doc, "Plan de acción priorizado")

    if not plan:
        p = doc.add_paragraph()
        r = p.add_run("No hay acciones pendientes: no se detectaron hábitos de riesgo.")
        r.font.name = FUENTE_TEXTO
        r.font.size = Pt(10.5)
        r.font.color.rgb = GOOD
        r.font.bold = True
    else:
        for i, item in enumerate(plan, start=1):
            color_sev = COLOR_RIESGO[ETIQUETA_SEVERIDAD[item["severidad"]]]

            p_item = doc.add_paragraph()
            p_item.paragraph_format.space_before = Pt(10)
            p_item.paragraph_format.space_after = Pt(1)
            r_num = p_item.add_run(f"{i}.  ")
            r_num.font.name = FUENTE_TITULO
            r_num.font.bold = True
            r_num.font.size = Pt(11)
            r_num.font.color.rgb = color_sev
            r_cat = p_item.add_run(f"[{NOMBRES_CATEGORIA[item['categoria']]}]  ")
            r_cat.font.name = FUENTE_TEXTO
            r_cat.font.bold = True
            r_cat.font.size = Pt(9)
            r_cat.font.color.rgb = INK_FAINT
            r_desc = p_item.add_run(item["descripcion"])
            r_desc.font.name = FUENTE_TEXTO
            r_desc.font.bold = True
            r_desc.font.size = Pt(10.5)
            r_desc.font.color.rgb = INK

            p_rec = doc.add_paragraph()
            p_rec.paragraph_format.left_indent = Inches(0.28)
            p_rec.paragraph_format.space_after = Pt(2)
            r_rec = p_rec.add_run(item["recomendacion"])
            r_rec.font.name = FUENTE_TEXTO
            r_rec.font.italic = True
            r_rec.font.size = Pt(10)
            r_rec.font.color.rgb = INK_SOFT

            p_imp = doc.add_paragraph()
            p_imp.paragraph_format.left_indent = Inches(0.28)
            p_imp.paragraph_format.space_after = Pt(4)
            r_imp = p_imp.add_run(
                f"{item['puntaje_actual']} → {item['puntaje_resultante']}   "
                f"(−{item['impacto']} pts · {ETIQUETA_ESFUERZO[item['esfuerzo']]})"
            )
            r_imp.font.name = FUENTE_DATOS
            r_imp.font.bold = True
            r_imp.font.size = Pt(9.5)
            r_imp.font.color.rgb = color_sev

    # --- Pie -----------------------------------------------------------------
    p_pie = doc.add_paragraph()
    p_pie.paragraph_format.space_before = Pt(24)
    r = p_pie.add_run(
        "Reporte generado localmente por Mapa de Riesgos Digitales. "
        "Ningún dato fue enviado a servidores externos."
    )
    r.font.name = FUENTE_TEXTO
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = INK_FAINT

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
